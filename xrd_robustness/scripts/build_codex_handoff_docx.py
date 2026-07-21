#!/usr/bin/env python3
"""Render CODEX_HANDOFF.md as a portable compact-reference DOCX."""

from __future__ import annotations

import argparse
from datetime import date
from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[1]
DEFAULT_SOURCE = PROJECT_ROOT / "CODEX_HANDOFF.md"
DEFAULT_OUTPUT = PROJECT_ROOT / "docs" / "CODEX_ACCOUNT_HANDOFF.docx"

BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
INK = RGBColor(0x0B, 0x25, 0x45)
MUTED = RGBColor(0x66, 0x66, 0x66)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CODE_GRAY = "F6F8FA"
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--source", default=str(DEFAULT_SOURCE))
    parser.add_argument("--output", default=str(DEFAULT_OUTPUT))
    return parser.parse_args()


def _set_run_font(
    run,
    *,
    name: str = "Calibri",
    east_asia: str = "Microsoft YaHei",
    size: float | None = None,
    color: RGBColor | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.insert(0, fonts)
    fonts.set(qn("w:ascii"), name)
    fonts.set(qn("w:hAnsi"), name)
    fonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def _set_cell_margins(cell, *, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_table_geometry(table, widths: list[int]) -> None:
    if sum(widths) != CONTENT_WIDTH_DXA:
        raise ValueError("table widths must sum to 9360 DXA")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            _set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def _add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    _set_run_font(run, size=9, color=MUTED)
    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")
    run._r.extend((field_begin, instr, field_end))


def _configure_document(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hr = hp.add_run("XRD Robustness V9-T  |  Cross-Codex Handoff")
    _set_run_font(hr, size=9, color=MUTED)
    footer = section.footer
    footer.is_linked_to_previous = False
    _add_page_field(footer.paragraphs[0])


def _add_inline(paragraph, text: str, *, base_size: float = 11) -> None:
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")
    position = 0
    for match in pattern.finditer(text):
        if match.start() > position:
            run = paragraph.add_run(text[position : match.start()])
            _set_run_font(run, size=base_size)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            _set_run_font(run, size=base_size, bold=True)
        else:
            run = paragraph.add_run(token[1:-1])
            _set_run_font(run, name="Consolas", east_asia="Microsoft YaHei", size=base_size - 0.5, color=INK)
            _shade_run(run, "EEF2F6")
        position = match.end()
    if position < len(text):
        run = paragraph.add_run(text[position:])
        _set_run_font(run, size=base_size)


def _shade_run(run, fill: str) -> None:
    rpr = run._element.get_or_add_rPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    rpr.append(shading)


def _add_code_block(doc: Document, lines: list[str]) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.15)
    paragraph.paragraph_format.right_indent = Inches(0.15)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.05
    ppr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), CODE_GRAY)
    ppr.append(shading)
    run = paragraph.add_run("\n".join(lines))
    _set_run_font(run, name="Consolas", east_asia="Microsoft YaHei", size=8.5, color=INK)


def _table_widths(column_count: int) -> list[int]:
    patterns = {
        2: [2700, 6660],
        3: [2100, 2000, 5260],
        4: [1500, 1300, 3280, 3280],
    }
    if column_count in patterns:
        return patterns[column_count]
    base = CONTENT_WIDTH_DXA // column_count
    widths = [base] * column_count
    widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
    return widths


def _add_markdown_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    column_count = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=column_count)
    table.style = "Table Grid"
    _set_table_geometry(table, _table_widths(column_count))
    for row_index, values in enumerate(rows):
        for column_index, value in enumerate(values):
            cell = table.cell(row_index, column_index)
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.15
            _add_inline(paragraph, value, base_size=9.2)
            if row_index == 0:
                _shade_cell(cell, LIGHT_BLUE)
                for run in paragraph.runs:
                    run.bold = True
            elif row_index % 2 == 0:
                _shade_cell(cell, LIGHT_GRAY)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def _parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    raw: list[list[str]] = []
    index = start
    while index < len(lines) and lines[index].strip().startswith("|"):
        values = [part.strip() for part in lines[index].strip().strip("|").split("|")]
        raw.append(values)
        index += 1
    if len(raw) >= 2 and all(re.fullmatch(r":?-{3,}:?", value) for value in raw[1]):
        raw.pop(1)
    return raw, index


def _add_cover(doc: Document) -> None:
    for _ in range(3):
        doc.add_paragraph()
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kr = kicker.add_run("CROSS-ACCOUNT OPERATIONS MANUAL")
    _set_run_font(kr, size=10, color=BLUE, bold=True)
    kicker.paragraph_format.space_after = Pt(18)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("XRD Robustness V9-T")
    _set_run_font(tr, size=28, color=INK, bold=True)
    title.paragraph_format.space_after = Pt(8)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = subtitle.add_run("跨 Codex 账号与台式机完整交接手册")
    _set_run_font(sr, size=17, color=DARK_BLUE, bold=True)
    subtitle.paragraph_format.space_after = Pt(12)

    description = doc.add_paragraph()
    description.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = description.add_run("研究边界 · 当前进度 · 迁移核验 · 硬件门禁 · 训练授权 · 故障恢复")
    _set_run_font(dr, size=11, color=MUTED)
    description.paragraph_format.space_after = Pt(48)

    status = doc.add_paragraph()
    status.alignment = WD_ALIGN_PARAGRAPH.CENTER
    status_run = status.add_run("状态：0/7 调参；台式机冷启动；当前无训练授权")
    _set_run_font(status_run, size=12, color=RGBColor(0x9B, 0x1C, 0x1C), bold=True)
    status.paragraph_format.space_after = Pt(8)

    snapshot = doc.add_paragraph()
    snapshot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    snap_run = snapshot.add_run(f"交接快照：{date.today().isoformat()}  |  权威入口：CODEX_HANDOFF.md")
    _set_run_font(snap_run, size=10, color=MUTED)
    doc.add_page_break()


def _add_contents(doc: Document, lines: list[str]) -> None:
    heading = doc.add_paragraph("目录", style="Heading 1")
    heading.paragraph_format.space_before = Pt(0)
    for line in lines:
        if line.startswith("## "):
            p = doc.add_paragraph(style="List Number")
            _add_inline(p, line[3:].strip(), base_size=10)
    doc.add_page_break()


def build(source: Path, output: Path) -> None:
    lines = source.read_text(encoding="utf-8").splitlines()
    doc = Document()
    _configure_document(doc)
    _add_cover(doc)
    _add_contents(doc, lines)

    index = 0
    in_code = False
    code_lines: list[str] = []
    skipped_source_title = False
    while index < len(lines):
        raw = lines[index]
        stripped = raw.strip()
        if stripped.startswith("```"):
            if in_code:
                _add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            index += 1
            continue
        if in_code:
            code_lines.append(raw)
            index += 1
            continue
        if not stripped:
            index += 1
            continue
        if stripped.startswith("|"):
            table_rows, index = _parse_table(lines, index)
            _add_markdown_table(doc, table_rows)
            continue
        if stripped.startswith("# ") and not skipped_source_title:
            skipped_source_title = True
            index += 1
            continue
        if stripped.startswith("#### "):
            p = doc.add_paragraph(stripped[5:], style="Heading 3")
            p.paragraph_format.left_indent = Inches(0.1)
        elif stripped.startswith("### "):
            doc.add_paragraph(stripped[4:], style="Heading 2")
        elif stripped.startswith("## "):
            doc.add_paragraph(stripped[3:], style="Heading 1")
        elif stripped.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.22)
            p.paragraph_format.right_indent = Inches(0.15)
            p.paragraph_format.space_after = Pt(8)
            ppr = p._p.get_or_add_pPr()
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), "F4F6F9")
            ppr.append(shading)
            _add_inline(p, stripped[2:], base_size=10.5)
            for run in p.runs:
                run.italic = True
        elif re.match(r"^- ", stripped):
            p = doc.add_paragraph(style="List Bullet")
            _add_inline(p, stripped[2:])
        elif re.match(r"^\d+\. ", stripped):
            p = doc.add_paragraph(style="List Number")
            _add_inline(p, re.sub(r"^\d+\. ", "", stripped, count=1))
        else:
            p = doc.add_paragraph()
            _add_inline(p, stripped)
        index += 1

    if in_code and code_lines:
        _add_code_block(doc, code_lines)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def main() -> int:
    args = parse_args()
    source = Path(args.source).resolve()
    output = Path(args.output).resolve()
    build(source, output)
    print(f"output={output}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
